"""
Генератор тестовых данных для table-merger.

Создаёт 5 xlsx и 5 docx файлов в текущей папке.

Все файлы — таблица «сотрудники компании», но с РАЗНЫМ набором столбцов
и РАЗНЫМ написанием одинаковых по смыслу заголовков. Это нагрузочный тест
для нормализатора заголовков и логики union столбцов:

  - Нормализация склеит:
      ФИО / Ф.И.О. / фио / Ф И О → один столбец
      Email / E-mail               → один столбец

  - Нормализация НЕ склеит (это разные слова — союз сохранит как разные):
      Дата приёма / Дата найма / Дата трудоустройства  → 3 столбца
      Стаж работы (лет) / Стаж, лет                    → 2 столбца

  - Union сохранит уникальные столбцы из отдельных файлов:
      Телефон, Email, Возраст, Образование, Бонус %, Стаж — каждый есть
      не во всех файлах, в строках без них итог будет пустой.

В файлах часть значений намеренно пустые (~10% ячеек), чтобы протестировать
обработку пропусков.

Запуск:
    cd C:\\Users\\Lenovo\\Desktop\\GigaChat\\test-files
    python generate-test-data.py

Требуются: openpyxl, python-docx
    pip install openpyxl python-docx
"""

import random
from datetime import date, timedelta
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document


# Фиксируем seed чтобы при повторном запуске получались те же файлы.
random.seed(42)


NAMES_M = [
    'Иванов Иван Иванович',
    'Петров Пётр Петрович',
    'Сидоров Алексей Николаевич',
    'Кузнецов Дмитрий Сергеевич',
    'Смирнов Владимир Андреевич',
    'Соколов Михаил Викторович',
    'Попов Андрей Дмитриевич',
    'Лебедев Сергей Александрович',
    'Козлов Артём Олегович',
    'Новиков Илья Романович',
    'Морозов Григорий Павлович',
    'Волков Антон Игоревич',
]
NAMES_F = [
    'Иванова Анна Сергеевна',
    'Петрова Мария Викторовна',
    'Сидорова Ольга Дмитриевна',
    'Кузнецова Екатерина Игоревна',
    'Смирнова Татьяна Николаевна',
    'Соколова Наталья Андреевна',
    'Попова Елена Александровна',
    'Лебедева Светлана Петровна',
    'Козлова Юлия Романовна',
    'Новикова Дарья Викторовна',
    'Морозова Анастасия Олеговна',
    'Волкова Полина Игоревна',
]

POSITIONS = [
    'Менеджер', 'Старший менеджер', 'Аналитик', 'Старший аналитик',
    'Разработчик', 'Старший разработчик', 'Тимлид', 'Дизайнер',
    'UX-исследователь', 'Бухгалтер', 'Финансовый аналитик',
    'Тестировщик', 'Специалист по продажам', 'HR-специалист',
    'Руководитель отдела',
]

DEPARTMENTS = [
    'Маркетинг', 'Разработка', 'Финансы',
    'Продажи', 'HR', 'Поддержка', 'Аналитика', 'Дизайн',
]

EDUCATION = [
    'Высшее техническое',
    'Высшее экономическое',
    'Высшее гуманитарное',
    'Магистратура',
    'Бакалавр',
    'Среднее специальное',
    'Незаконченное высшее',
]


# Транслитерация фамилии для email. Простая ASCII-table.
_TRANSLIT = {
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e',
    'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
    'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
    'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sch',
    'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
}


def translit(s: str) -> str:
    return ''.join(_TRANSLIT.get(c.lower(), c.lower()) for c in s if c.isalpha())


# Описание файлов: (имя, заголовки).
# Намеренно ОЧЕНЬ разные структуры — от минимума (5 столбцов) до полного
# набора (9 столбцов). Разные написания одного и того же столбца.
FILE_VARIATIONS = [
    # Минимальный — только базовые поля
    ('маркетинг',  [
        'ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
    ]),
    # Базовый + контакты + стаж, ФИО через точки
    ('разработка', [
        'Ф.И.О.', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
        'Email', 'Телефон', 'Стаж работы (лет)',
    ]),
    # Все заголовки в нижнем регистре + образование + ИНАЯ КОЛОНКА «Дата найма»
    ('финансы',    [
        'фио', 'должность', 'отдел', 'зарплата', 'дата найма',
        'образование',
    ]),
    # Самый широкий — 9 столбцов, с возрастом и бонусом
    ('продажи',    [
        'ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
        'Телефон', 'Email', 'Возраст', 'Бонус, %',
    ]),
    # ФИО через пробелы + ЕЩЁ ОДНА вариация даты + email с дефисом + стаж по-другому
    ('hr',         [
        'Ф И О', 'Должность', 'Отдел', 'Зарплата', 'Дата трудоустройства',
        'E-mail', 'Стаж, лет',
    ]),
]


def normalize_for_match(h: str) -> str:
    """Тот же нормализатор что в браузерном table-merger.
    Используем здесь чтобы понимать какой смысл несёт каждый заголовок."""
    import re
    return re.sub(r"[\s._\-(){}\[\]/\\:;,!?\"' ]+", '', str(h).strip().lower())


def gen_value_for_header(header: str, name_for_email: str = '') -> object:
    """По имени заголовка генерируем подходящее значение.
    name_for_email — фамилия чтобы сделать осмысленный email."""
    key = normalize_for_match(header)

    # ~10% значений делаем пустыми (NULL) — для теста пропусков.
    # Кроме обязательных полей (ФИО, Должность, Отдел).
    is_optional = not any(k in key for k in ('фио', 'должность', 'отдел'))
    if is_optional and random.random() < 0.1:
        return ''

    if 'фио' in key:
        return random.choice(NAMES_M + NAMES_F)
    if 'должность' in key:
        return random.choice(POSITIONS)
    if 'отдел' in key:
        return random.choice(DEPARTMENTS)
    if 'зарплата' in key:
        # Иногда круглые тысячи, иногда копейки
        if random.random() < 0.3:
            return round(random.uniform(50_000, 300_000), 2)
        return random.randint(50_000, 300_000)
    # Все три варианта даты («дата приёма», «дата найма», «дата трудоустройства»)
    # генерируем одинаково — это просто разные слова для одного смысла.
    if 'дата' in key:
        d = date(2018, 1, 1) + timedelta(days=random.randint(0, 365 * 7))
        return d.strftime('%d.%m.%Y')
    if 'телефон' in key:
        return '+7' + ''.join(str(random.randint(0, 9)) for _ in range(10))
    if 'email' in key or 'mail' in key:
        # фамилия@company.ru — translit фамилии
        surname = (name_for_email or '').split()[0] if name_for_email else 'user'
        return f'{translit(surname)}@gigachat-corp.ru'
    if 'возраст' in key:
        return random.randint(22, 62)
    if 'образование' in key:
        return random.choice(EDUCATION)
    if 'бонус' in key:
        return random.randint(0, 30)
    if 'стаж' in key:
        return random.randint(0, 18)
    return '?'


def gen_row(headers: list) -> list:
    """Генерируем одну строку, передавая имя дальше в gen для email."""
    # Сначала ФИО (для использования в email)
    name = ''
    for h in headers:
        if 'фио' in normalize_for_match(h):
            name = random.choice(NAMES_M + NAMES_F)
            break

    row = []
    used_name = False
    for h in headers:
        if 'фио' in normalize_for_match(h):
            if not used_name and name:
                row.append(name)
                used_name = True
            else:
                row.append(random.choice(NAMES_M + NAMES_F))
        else:
            row.append(gen_value_for_header(h, name))
    return row


# ---------- XLSX ----------

def make_xlsx(filepath: Path, headers: list, n_rows: int) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = 'Сотрудники'

    header_font = Font(bold=True, color='FFFFFFFF')
    header_fill = PatternFill('solid', fgColor='FF4F46E5')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)

    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    for r in range(n_rows):
        row_data = gen_row(headers)
        for col_idx, val in enumerate(row_data, start=1):
            ws.cell(row=r + 2, column=col_idx, value=val)

    # Авто-ширина (грубая)
    for col_idx, h in enumerate(headers, start=1):
        max_len = max(len(str(h)), 12)
        for r in range(n_rows):
            v = ws.cell(row=r + 2, column=col_idx).value
            if v is not None and len(str(v)) > max_len:
                max_len = len(str(v))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_len + 2, 35)

    ws.freeze_panes = 'A2'
    wb.save(filepath)


# ---------- DOCX ----------

def make_docx(filepath: Path, headers: list, n_rows: int, dept_label: str) -> None:
    doc = Document()
    doc.add_heading(f'Отчёт по отделу: {dept_label}', level=1)
    doc.add_paragraph(
        'Источник: HR-отдел. Период: 2024 год. '
        'Документ сгенерирован автоматически для тестирования table-merger.'
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for r in range(n_rows):
        row_data = gen_row(headers)
        row_cells = table.rows[r + 1].cells
        for col_idx, val in enumerate(row_data):
            row_cells[col_idx].text = '' if val is None or val == '' else str(val)

    doc.add_paragraph()
    doc.add_paragraph(
        'Этот текст идёт ПОСЛЕ таблицы — он также должен быть проигнорирован.'
    )
    doc.save(filepath)


# ---------- main ----------

def main() -> None:
    out_dir = Path(__file__).parent
    print(f'Генерирую файлы в: {out_dir}')
    print()

    summary = []
    for idx, (name, headers) in enumerate(FILE_VARIATIONS, start=1):
        n_rows_x = random.randint(5, 15)
        xlsx_path = out_dir / f'{name}.xlsx'
        make_xlsx(xlsx_path, headers, n_rows_x)

        n_rows_d = random.randint(5, 15)
        docx_path = out_dir / f'{name}.docx'
        make_docx(docx_path, headers, n_rows_d, name.capitalize())

        summary.append((name, headers, n_rows_x, n_rows_d))
        print(f'  [{idx}/{len(FILE_VARIATIONS)}] {name}: '
              f'{len(headers)} колонок | xlsx={n_rows_x} строк | docx={n_rows_d} строк')

    print()
    print('Готово.')
    print()
    print('Сводка по заголовкам (видно вариативность для теста):')
    for name, headers, _, _ in summary:
        print(f'  {name}: {headers}')

    print()
    print('Что должно произойти при объединении ВСЕХ 5 файлов одного формата:')
    print('  ──── Слепливается в одну колонку ────')
    print('   ФИО (4 формы: ФИО, Ф.И.О., фио, Ф И О) → 1 колонка')
    print('   Email/E-mail → 1 колонка')
    print('   Зарплата (5 файлов) → 1 колонка')
    print('   Должность, Отдел → по 1 колонке')
    print('  ──── Останется как разные колонки (намеренно) ────')
    print('   Дата приёма / Дата найма / Дата трудоустройства → 3 колонки')
    print('   Стаж работы (лет) / Стаж, лет → 2 колонки')
    print('  ──── Уникальные (есть только в части файлов) ────')
    print('   Телефон, Возраст, Образование, Бонус — в строках без этих')
    print('   полей в итоге будет пусто.')
    print('  ──── Пропуски ────')
    print('   ~10% ячеек намеренно пустые — проверка обработки NULL.')


if __name__ == '__main__':
    main()
