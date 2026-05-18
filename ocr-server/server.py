"""
Локальный HTTP-сервис извлечения текста из документов для GigaChat.

Делает то же, что коммерческий OCR-as-a-service, но полностью офлайн:
  - PDF: сначала пробует прямое извлечение через PyMuPDF (быстро, точно).
         Если в PDF мало текста — рендерит страницы как картинки и пропускает
         через EasyOCR (для сканированных документов).
  - DOCX: прямое извлечение через python-docx.
  - Картинки (PNG/JPG/TIFF/BMP/WEBP): EasyOCR.
  - TXT/MD/CSV: просто читаем как UTF-8.

n8n-workflow «organization-appeal» (и другие, использующие OCR) шлют файлы
на эндпоинт POST /v1/file/text/ с полем multipart `file` и ожидают JSON
с полем `text`.

Запуск:
    pip install -r requirements.txt
    python server.py

Конфигурация через переменные окружения:
    OCR_HOST           — на каком интерфейсе слушать. По умолчанию: 0.0.0.0
    OCR_PORT           — порт. По умолчанию: 8080
    OCR_LANGS          — языки для EasyOCR. По умолчанию: ru,en
    OCR_PDF_MIN_TEXT   — если PDF после PyMuPDF дал меньше N символов —
                          fallback на EasyOCR. По умолчанию: 50
    OCR_PDF_MAX_PAGES  — лимит страниц PDF для OCR (защита от мега-документов).
                          По умолчанию: 50
    OCR_EASYOCR_DIR    — папка с моделями EasyOCR (см. README). Если не задано —
                          ~/.EasyOCR (стандартное место).
    OCR_DEVICE         — cpu / cuda. По умолчанию: auto.

API:
    GET  /health
        resp: {"status": "ok", "easyocr_ready": bool, "device": "cpu|cuda"}

    POST /v1/file/text/
        multipart: file=<binary>
        resp: {"text": "...", "source": "pymupdf|pymupdf+easyocr|docx|easyocr|plain", "filename": "..."}

    POST /v1/image/text/
        multipart: file=<binary>
        resp: {"text": "...", "source": "easyocr"}
"""

import io
import os
import logging
import tempfile
from typing import Optional

# Принудительный оффлайн-режим: EasyOCR/torch не должны лезть в интернет
# даже если detect что моделей нет. Безопаснее, чем дефолт.
os.environ.setdefault('HF_HUB_OFFLINE', '1')
os.environ.setdefault('TRANSFORMERS_OFFLINE', '1')

import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException

HOST = os.environ.get('OCR_HOST', '0.0.0.0')
PORT = int(os.environ.get('OCR_PORT', '8080'))
LANGS = os.environ.get('OCR_LANGS', 'ru,en').split(',')
PDF_MIN_TEXT = int(os.environ.get('OCR_PDF_MIN_TEXT', '50'))
PDF_MAX_PAGES = int(os.environ.get('OCR_PDF_MAX_PAGES', '50'))
EASYOCR_DIR = os.environ.get('OCR_EASYOCR_DIR', '').strip() or None
DEVICE = os.environ.get('OCR_DEVICE', '').strip().lower()

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
)
log = logging.getLogger('ocr-server')


# --- Ленивая инициализация EasyOCR ---
# EasyOCR грузит модели на ~150 МБ — делаем это при первом обращении,
# чтобы /health отвечал моментально (workflow его пингует часто).
_easyocr_reader = None
_easyocr_init_error: Optional[str] = None


def _detect_device() -> str:
    if DEVICE in ('cpu', 'cuda'):
        return DEVICE
    try:
        import torch
        if torch.cuda.is_available():
            return 'cuda'
    except Exception:
        pass
    return 'cpu'


def get_easyocr():
    """Инициализирует EasyOCR при первом вызове; повторно — возвращает кеш."""
    global _easyocr_reader, _easyocr_init_error
    if _easyocr_reader is not None:
        return _easyocr_reader
    if _easyocr_init_error is not None:
        # Не пытаемся повторно — если упало один раз, упадёт снова.
        raise RuntimeError(_easyocr_init_error)
    try:
        import easyocr
        use_gpu = _detect_device() == 'cuda'
        log.info('Loading EasyOCR (langs=%s, gpu=%s, dir=%s)', LANGS, use_gpu, EASYOCR_DIR or 'default')
        _easyocr_reader = easyocr.Reader(
            LANGS,
            gpu=use_gpu,
            model_storage_directory=EASYOCR_DIR,
            download_enabled=False,  # КРИТИЧНО: офлайн. Если моделей нет — упадёт явно.
            verbose=False,
        )
        log.info('EasyOCR ready.')
        return _easyocr_reader
    except Exception as e:
        _easyocr_init_error = (
            f'EasyOCR недоступен: {e}. '
            f'Проверь, что модели лежат в {EASYOCR_DIR or "~/.EasyOCR/model/"} '
            'и EASYOCR_DIR указан корректно.'
        )
        log.error(_easyocr_init_error)
        raise RuntimeError(_easyocr_init_error)


# --- Извлечение текста по типам ---

def extract_pdf(data: bytes) -> tuple[str, str]:
    """
    PDF: сначала PyMuPDF (прямое извлечение). Если получено мало текста —
    рендерим страницы как картинки и пропускаем через EasyOCR.
    Возвращает (text, source).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype='pdf')
    try:
        if doc.page_count > PDF_MAX_PAGES:
            log.warning('PDF has %d pages, limiting to %d', doc.page_count, PDF_MAX_PAGES)

        # Шаг 1: прямое извлечение через PyMuPDF
        parts = []
        for i, page in enumerate(doc):
            if i >= PDF_MAX_PAGES:
                break
            parts.append(page.get_text('text'))
        text = '\n'.join(parts).strip()

        if len(text) >= PDF_MIN_TEXT:
            return text, 'pymupdf'

        # Шаг 2: PDF без текстового слоя — рендерим страницы и OCR-им
        log.info('PDF text too short (%d chars), falling back to EasyOCR', len(text))
        reader = get_easyocr()
        ocr_parts = []
        for i, page in enumerate(doc):
            if i >= PDF_MAX_PAGES:
                break
            # 2x — увеличенное разрешение для лучшего распознавания
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes('png')
            # EasyOCR принимает numpy-array или путь к файлу
            import numpy as np
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            arr = np.array(img)
            result = reader.readtext(arr, detail=0, paragraph=True)
            ocr_parts.append('\n'.join(result))
        full = '\n'.join(ocr_parts).strip()
        return full, 'pymupdf+easyocr'
    finally:
        doc.close()


def extract_docx(data: bytes) -> tuple[str, str]:
    """DOCX: через python-docx (читаем параграфы + таблицы)."""
    import docx
    f = io.BytesIO(data)
    doc = docx.Document(f)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = '\n'.join(parts).strip()
    return text, 'docx'


def extract_image(data: bytes) -> tuple[str, str]:
    """Картинки → EasyOCR."""
    import numpy as np
    from PIL import Image
    reader = get_easyocr()
    img = Image.open(io.BytesIO(data))
    # Гарантируем RGB (EasyOCR не любит RGBA)
    if img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')
    arr = np.array(img)
    result = reader.readtext(arr, detail=0, paragraph=True)
    text = '\n'.join(result).strip()
    return text, 'easyocr'


def extract_plain(data: bytes) -> tuple[str, str]:
    """TXT/MD/CSV: пробуем UTF-8, потом cp1251 (распространено в офисных файлах)."""
    for enc in ('utf-8', 'utf-8-sig', 'cp1251', 'latin-1'):
        try:
            return data.decode(enc).strip(), 'plain'
        except UnicodeDecodeError:
            continue
    raise HTTPException(415, 'Не удалось декодировать текстовый файл')


# --- FastAPI ---

app = FastAPI(title='GigaChat OCR Server')


@app.get('/health')
def health():
    """Лёгкий health-check. НЕ инициализирует EasyOCR."""
    return {
        'status': 'ok',
        'easyocr_ready': _easyocr_reader is not None,
        'easyocr_error': _easyocr_init_error,
        'device': _detect_device(),
        'langs': LANGS,
        'port': PORT,
    }


IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'tiff', 'tif', 'bmp', 'webp'}
PLAIN_EXTS = {'txt', 'md', 'csv', 'log'}


@app.post('/v1/file/text/')
async def file_text(file: UploadFile = File(...)):
    """
    Основной эндпоинт. Принимает любой документ, отдаёт извлечённый текст.
    Совместим с тем, что ожидают workflow'ы organization-appeal, document-loader и др.
    """
    if not file.filename:
        raise HTTPException(400, 'Файл без имени')
    data = await file.read()
    if not data:
        raise HTTPException(400, 'Пустой файл')

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    log.info('Received: %s (%d bytes, ext=%s)', file.filename, len(data), ext)

    try:
        if ext == 'pdf':
            text, source = extract_pdf(data)
        elif ext == 'docx':
            text, source = extract_docx(data)
        elif ext in IMAGE_EXTS:
            text, source = extract_image(data)
        elif ext in PLAIN_EXTS:
            text, source = extract_plain(data)
        else:
            raise HTTPException(415, f'Неподдерживаемый формат: .{ext}. Поддерживаются: pdf, docx, txt, png, jpg, tiff, bmp, webp.')
    except HTTPException:
        raise
    except RuntimeError as e:
        # EasyOCR не загрузился — отдадим понятную ошибку
        raise HTTPException(503, str(e))
    except Exception as e:
        log.exception('Extract error')
        raise HTTPException(500, f'Ошибка извлечения: {e}')

    log.info('Extracted %d chars via %s', len(text), source)
    return {
        'text': text,
        'source': source,
        'filename': file.filename,
        'chars': len(text),
    }


@app.post('/v1/image/text/')
async def image_text(file: UploadFile = File(...)):
    """Отдельный эндпоинт для картинок — всегда через EasyOCR."""
    data = await file.read()
    if not data:
        raise HTTPException(400, 'Пустой файл')
    try:
        text, source = extract_image(data)
    except RuntimeError as e:
        raise HTTPException(503, str(e))
    except Exception as e:
        log.exception('Image OCR error')
        raise HTTPException(500, str(e))
    return {'text': text, 'source': source, 'filename': file.filename, 'chars': len(text)}


if __name__ == '__main__':
    log.info('Starting OCR server on http://%s:%d', HOST, PORT)
    log.info('Langs: %s | PDF min text: %d | PDF max pages: %d',
             LANGS, PDF_MIN_TEXT, PDF_MAX_PAGES)
    uvicorn.run(app, host=HOST, port=PORT, log_level='info')
