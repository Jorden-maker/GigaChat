"""
Локальный HTTP-сервис объединения таблиц для GigaChat.

Принимает несколько файлов одного формата (xlsx или docx), извлекает из каждого
одну таблицу, объединяет в одну общую и возвращает итоговый файл того же формата.

Заголовки столбцов сопоставляются по нормализованному виду (без пробелов,
точек, скобок, регистра), так что «Ф.И.О.», «Ф И О» и «ФИО» считаются одним
столбцом. Если в одном файле есть столбец, которого нет в других — он
добавляется в итог, остальные строки в нём пустые.

Запуск:
    pip install -r requirements.txt
    python server.py

Конфигурация через переменные окружения:
    MERGER_HOST  — на каком интерфейсе слушать. По умолчанию 0.0.0.0
    MERGER_PORT  — порт. По умолчанию 8082
    MERGER_MAX_FILES — жёсткий лимит на число файлов. По умолчанию 25.

API:
    POST /merge
        multipart/form-data, поле files[] (несколько файлов)
        Все файлы должны быть одного формата (.xlsx или .docx).
        Возвращает binary с MIME-type и Content-Disposition.

    GET /health
        Проверка живости сервиса.
"""

import io
import os
import re
import logging
from typing import List, Tuple, Dict, Optional
from urllib.parse import quote as url_quote

import uvicorn
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware

from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document


HOST = os.environ.get('MERGER_HOST', '0.0.0.0')
PORT = int(os.environ.get('MERGER_PORT', '8082'))
MAX_FILES = int(os.environ.get('MERGER_MAX_FILES', '25'))
WARN_FILES = 10  # порог предупреждения в UI

XLSX_MIME = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
)
log = logging.getLogger('table-merger')


# ---------------------------------------------------------------------------
# Нормализация заголовков
# ---------------------------------------------------------------------------

_NORM_RE = re.compile(r'[\s._\-(){}\[\]/\\\:;,!?\"\' ]+')


def normalize_header(h) -> str:
    """«Ф.И.О.» / «Ф И О» / «фио» → «фио». Используется для сопоставления столбцов."""
    if h is None:
        return ''
    s = str(h).strip().lower()
    s = _NORM_RE.sub('', s)
    return s


def cell_to_text(v) -> str:
    """Значение ячейки → текст для записи. Без округления."""
    if v is None:
        return ''
    if isinstance(v, bool):
        return 'TRUE' if v else 'FALSE'
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


# ---------------------------------------------------------------------------
# Чтение таблицы из файла
# ---------------------------------------------------------------------------

def read_xlsx_table(data: bytes, filename: str) -> Tuple[List[str], List[List]]:
    """Читаем первый лист xlsx, считаем первую строку заголовком."""
    try:
        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    except Exception as e:
        raise HTTPException(400, f"Не удалось открыть Excel-файл «{filename}»: {e}")

    ws = wb.worksheets[0]
    rows_iter = ws.iter_rows(values_only=True)
    header_row = next(rows_iter, None)
    if header_row is None:
        wb.close()
        raise HTTPException(400, f"Файл «{filename}» пустой (нет данных на первом листе).")

    # Обрезаем хвостовые пустые ячейки в заголовке (типичный артефакт Excel)
    headers = list(header_row)
    while headers and (headers[-1] is None or str(headers[-1]).strip() == ''):
        headers.pop()
    if not headers:
        wb.close()
        raise HTTPException(400, f"В файле «{filename}» первая строка пустая — нет заголовков.")
    headers = [str(h).strip() if h is not None else '' for h in headers]

    rows: List[List] = []
    for row in rows_iter:
        # Пропускаем полностью пустые строки
        if not any(v is not None and str(v).strip() != '' for v in row):
            continue
        # Обрезаем до длины заголовков, добиваем пустыми если короче
        clean = list(row[:len(headers)])
        while len(clean) < len(headers):
            clean.append(None)
        rows.append(clean)

    wb.close()
    return headers, rows


def read_docx_table(data: bytes, filename: str) -> Tuple[List[str], List[List]]:
    """Читаем первую таблицу в документе. Первая строка таблицы — заголовок."""
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(400, f"Не удалось открыть Word-файл «{filename}»: {e}")

    if not doc.tables:
        raise HTTPException(400, f"В файле «{filename}» нет таблиц.")

    table = doc.tables[0]
    if len(table.rows) < 1:
        raise HTTPException(400, f"Таблица в файле «{filename}» пустая.")

    headers = [cell.text.strip() for cell in table.rows[0].cells]
    # Word допускает merged-ячейки — может появиться дубль одного и того же объекта.
    # Дедупликация по object id, чтобы в headers были уникальные ячейки.
    seen_ids = set()
    unique_headers = []
    for h, cell in zip(headers, table.rows[0].cells):
        cid = id(cell._tc)
        if cid in seen_ids:
            continue
        seen_ids.add(cid)
        unique_headers.append(h)
    headers = unique_headers
    while headers and headers[-1] == '':
        headers.pop()
    if not headers:
        raise HTTPException(400, f"В файле «{filename}» в таблице нет заголовков.")

    rows: List[List] = []
    for tr in table.rows[1:]:
        seen_ids = set()
        cells = []
        for cell in tr.cells:
            cid = id(cell._tc)
            if cid in seen_ids:
                continue
            seen_ids.add(cid)
            cells.append(cell.text.strip())
        if not any(c != '' for c in cells):
            continue
        # Обрезаем/добиваем до длины headers
        cells = cells[:len(headers)]
        while len(cells) < len(headers):
            cells.append('')
        rows.append(cells)

    return headers, rows


# ---------------------------------------------------------------------------
# Слияние
# ---------------------------------------------------------------------------

def merge_tables(
    sources: List[Tuple[str, List[str], List[List]]],
) -> Tuple[List[str], List[List]]:
    """
    sources: список (filename, headers, rows) — по одному элементу на файл.
    Возвращает (объединённые заголовки, объединённые строки).

    Алгоритм:
      1. По всем файлам собираем список нормализованных столбцов, сохраняя
         порядок первого появления (стабильность для пользователя).
      2. Имя столбца в итоге = первое встреченное оригинальное имя.
      3. Для каждой строки из каждого файла строим новую строку с правильным
         порядком столбцов. Если в исходной таблице столбца не было — пусто.
    """
    column_order: List[str] = []                 # нормализованные ключи в порядке появления
    canonical: Dict[str, str] = {}               # норм. ключ -> отображаемое имя

    for filename, headers, _rows in sources:
        for h in headers:
            key = normalize_header(h)
            if not key:
                # Пустой заголовок встречается — игнорируем, чтобы не было
                # «безымянного» столбца, который собирает мусор от разных файлов.
                continue
            if key not in canonical:
                canonical[key] = h
                column_order.append(key)

    if not column_order:
        raise HTTPException(400, "Во всех файлах заголовки пустые — нечего объединять.")

    merged_headers = [canonical[k] for k in column_order]
    merged_rows: List[List] = []

    for filename, headers, rows in sources:
        # Для текущего файла строим маппинг норм.ключ -> индекс столбца в исходной таблице
        key_to_idx: Dict[str, int] = {}
        for idx, h in enumerate(headers):
            key = normalize_header(h)
            if key and key not in key_to_idx:
                key_to_idx[key] = idx

        for row in rows:
            new_row = []
            for key in column_order:
                src_idx = key_to_idx.get(key)
                if src_idx is None or src_idx >= len(row):
                    new_row.append('')
                else:
                    new_row.append(row[src_idx])
            merged_rows.append(new_row)

    return merged_headers, merged_rows


# ---------------------------------------------------------------------------
# Запись результата
# ---------------------------------------------------------------------------

def write_xlsx(headers: List[str], rows: List[List]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = 'Объединённая таблица'

    header_font = Font(bold=True, color='FFFFFFFF')
    header_fill = PatternFill('solid', fgColor='FF4F46E5')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)

    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    for row_idx, row in enumerate(rows, start=2):
        for col_idx, value in enumerate(row, start=1):
            ws.cell(row=row_idx, column=col_idx, value=cell_to_text(value) if value is not None else '')

    # Авто-ширина: считаем максимум по столбцу, ограничиваем разумным потолком
    for col_idx, h in enumerate(headers, start=1):
        max_len = len(str(h))
        for row in rows:
            v = row[col_idx - 1] if col_idx - 1 < len(row) else ''
            text = cell_to_text(v) if v is not None else ''
            if len(text) > max_len:
                max_len = len(text)
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max(max_len + 2, 10), 50)

    ws.freeze_panes = 'A2'

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


def write_docx(headers: List[str], rows: List[List]) -> bytes:
    doc = Document()
    doc.add_heading('Объединённая таблица', level=1)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Заголовок
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Данные
    for row_idx, row in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row):
            row_cells[col_idx].text = cell_to_text(value) if value is not None else ''

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# FastAPI
# ---------------------------------------------------------------------------

app = FastAPI(title='GigaChat Table Merger')

# CORS: UI может дёргать сервис как напрямую (для скачивания файла), так и
# через n8n-proxy. В офисной LAN это не open-internet, поэтому "*" безопасно.
app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],
    allow_credentials=False,
    allow_methods=['*'],
    allow_headers=['*'],
    expose_headers=['Content-Disposition'],
)


@app.get('/health')
def health():
    return {'status': 'ok', 'port': PORT, 'max_files': MAX_FILES}


def detect_kind(filename: str) -> Optional[str]:
    lower = filename.lower()
    if lower.endswith('.xlsx') or lower.endswith('.xlsm'):
        return 'xlsx'
    if lower.endswith('.docx'):
        return 'docx'
    return None


@app.post('/merge')
async def merge(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, 'Не передано ни одного файла.')
    if len(files) > MAX_FILES:
        raise HTTPException(
            400,
            f'Слишком много файлов: {len(files)} (предел {MAX_FILES}). '
            f'Объедини в несколько подходов.',
        )

    # Определяем формат — все файлы должны быть одного типа
    kinds = set()
    for f in files:
        k = detect_kind(f.filename or '')
        if k is None:
            raise HTTPException(
                400,
                f"Файл «{f.filename}» имеет неподдерживаемое расширение. "
                f"Поддерживаются только .xlsx, .xlsm и .docx.",
            )
        kinds.add(k)

    if len(kinds) > 1:
        raise HTTPException(
            400,
            'В одном запросе нельзя смешивать Excel и Word — '
            'выбери один формат для всей пачки.',
        )

    kind = kinds.pop()
    log.info('Merge: %d files, kind=%s', len(files), kind)

    # Читаем все файлы в память и парсим таблицы
    sources: List[Tuple[str, List[str], List[List]]] = []
    for f in files:
        data = await f.read()
        if not data:
            raise HTTPException(400, f"Файл «{f.filename}» пустой.")
        if kind == 'xlsx':
            headers, rows = read_xlsx_table(data, f.filename or 'без имени')
        else:
            headers, rows = read_docx_table(data, f.filename or 'без имени')
        sources.append((f.filename or '', headers, rows))
        log.info('  %s: %d cols, %d rows', f.filename, len(headers), len(rows))

    merged_headers, merged_rows = merge_tables(sources)
    log.info('Merged: %d cols, %d rows', len(merged_headers), len(merged_rows))

    # Пишем результат в нужном формате
    if kind == 'xlsx':
        out_bytes = write_xlsx(merged_headers, merged_rows)
        out_filename = 'merged.xlsx'
        mime = XLSX_MIME
    else:
        out_bytes = write_docx(merged_headers, merged_rows)
        out_filename = 'merged.docx'
        mime = DOCX_MIME

    # Content-Disposition c filename* для UTF-8 (на случай если когда-нибудь
    # выходное имя будет с кириллицей; сейчас latin, но привычка хорошая).
    cd = f"attachment; filename={out_filename}; filename*=UTF-8''{url_quote(out_filename)}"
    return Response(content=out_bytes, media_type=mime, headers={'Content-Disposition': cd})


if __name__ == '__main__':
    log.info('Starting on http://%s:%d (max_files=%d)', HOST, PORT, MAX_FILES)
    uvicorn.run(app, host=HOST, port=PORT, log_level='info')
